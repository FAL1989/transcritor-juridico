from __future__ import annotations

import io
import re
from typing import Dict, Optional

# Lazy imports for optional features (avoid crashing app if packages are missing)


def normalize_text(text: str) -> str:
    """Apply simple normalization suitable for legal documents.

    - Normalize line endings
    - Collapse multiple spaces
    - Ensure spacing around punctuation
    """
    if not text:
        return ""
    # Normalize line endings
    normalized = text.replace("\r\n", "\n").replace("\r", "\n")
    # Collapse excessive spaces but preserve newlines
    normalized = re.sub(r"[ \t]{2,}", " ", normalized)
    # Trim trailing spaces per line
    normalized = "\n".join(line.rstrip() for line in normalized.split("\n"))
    return normalized.strip()


def anonymize_text(text: str) -> str:
    """Optionally mask sensitive patterns (CPF/CNPJ, emails). Minimal MVP.

    Note: do not attempt to mask names automatically to avoid false positives.
    """
    if not text:
        return ""
    masked = text
    # Emails
    masked = re.sub(r"([A-Za-z0-9._%+-])[A-Za-z0-9._%+-]*(@[A-Za-z0-9.-]+\.[A-Za-z]{2,})", r"\1***\2", masked)
    # CPF: 000.000.000-00
    masked = re.sub(r"\b(\d{3})\.(\d{3})\.(\d{3})-(\d{2})\b", r"\1.***.***-**", masked)
    # CNPJ: 00.000.000/0000-00
    masked = re.sub(r"\b(\d{2})\.(\d{3})\.(\d{3})/(\d{4})-(\d{2})\b", r"\1.***.***/****-**", masked)
    return masked


def apply_template(template: Optional[str], content: str, metadata: Optional[Dict[str, str]] = None) -> str:
    """Apply a simple textual template. Supported: 'termo_audiencia', 'despacho'."""
    if not template:
        return content
    md = metadata or {}
    header_common = (
        f"{md.get('tribunal', 'Poder Judiciário')}\n"
        f"{md.get('comarca', '')}\n"
        f"Processo: {md.get('processo', 'N/D')}\n\n"
    )
    if template == "termo_audiencia":
        title = "TERMO DE AUDIÊNCIA"
        body = content
        return f"{header_common}{title}\n\n{body}\n\nLocal e data: {md.get('local_data', '')}\nAssinaturas: ______________________\n"
    if template == "despacho":
        title = "DESPACHO"
        body = content
        return f"{header_common}{title}\n\n{body}\n\nPublique-se. Intime-se. Cumpra-se.\n"
    return content


def extract_text_from_pdf(file_bytes: bytes) -> str:
    try:
        from PyPDF2 import PdfReader  # type: ignore
    except Exception:  # pragma: no cover
        return ""
    with io.BytesIO(file_bytes) as bio:
        reader = PdfReader(bio)
        parts = []
        for page in reader.pages:
            try:
                txt = page.extract_text() or ""
            except Exception:
                txt = ""
            parts.append(txt)
        return "\n\n".join(parts).strip()


def extract_text_from_docx(file_bytes: bytes) -> str:
    try:
        from docx import Document  # type: ignore
    except Exception:  # pragma: no cover
        return ""
    with io.BytesIO(file_bytes) as bio:
        doc = Document(bio)
        paragraphs = [p.text for p in doc.paragraphs]
        return "\n".join(paragraphs).strip()


def generate_docx_bytes(title: Optional[str], content: str, metadata: Optional[Dict[str, str]] = None) -> bytes:
    try:
        from docx import Document  # type: ignore
    except Exception:  # pragma: no cover
        # Return plain bytes as fallback
        return content.encode("utf-8")
    doc = Document()
    if title:
        doc.add_heading(title, level=1)
    if metadata:
        meta_lines = []
        if metadata.get("tribunal"):
            meta_lines.append(metadata["tribunal"])
        if metadata.get("comarca"):
            meta_lines.append(metadata["comarca"])
        if metadata.get("processo"):
            meta_lines.append(f"Processo: {metadata['processo']}")
        if meta_lines:
            doc.add_paragraph("\n".join(meta_lines))
            doc.add_paragraph("")
    for line in content.split("\n"):
        doc.add_paragraph(line)
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.read()


